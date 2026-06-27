from pathlib import Path
from typing import List

import pandas as pd
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader


def load_pdf(file_path: Path) -> List[Document]:
    loader = PyPDFLoader(str(file_path))
    documents = loader.load()

    for doc in documents:
        doc.metadata["source"] = file_path.name
        doc.metadata["file_type"] = "pdf"

    return documents


def load_txt(file_path: Path) -> List[Document]:
    loader = TextLoader(str(file_path), encoding="utf-8")
    documents = loader.load()

    for doc in documents:
        doc.metadata["source"] = file_path.name
        doc.metadata["file_type"] = "txt"
        doc.metadata["page"] = 0

    return documents


def load_docx(file_path: Path) -> List[Document]:
    docx_file = DocxDocument(str(file_path))

    text = "\n".join(
        paragraph.text
        for paragraph in docx_file.paragraphs
        if paragraph.text.strip()
    )

    return [
        Document(
            page_content=text,
            metadata={
                "source": file_path.name,
                "file_type": "docx",
                "page": 0
            }
        )
    ]


def load_csv(file_path: Path) -> List[Document]:
    df = pd.read_csv(file_path)

    documents = []

    for index, row in df.iterrows():
        row_text = "\n".join(
            f"{column}: {row[column]}"
            for column in df.columns
        )

        documents.append(
            Document(
                page_content=row_text,
                metadata={
                    "source": file_path.name,
                    "file_type": "csv",
                    "row": int(index),
                    "page": 0
                }
            )
        )

    return documents


def load_single_document(file_path: Path) -> List[Document]:
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return load_pdf(file_path)

    if suffix == ".txt":
        return load_txt(file_path)

    if suffix == ".docx":
        return load_docx(file_path)

    if suffix == ".csv":
        return load_csv(file_path)

    raise ValueError(f"Unsupported file type: {suffix}")


def load_documents_from_folder(folder_path: Path) -> List[Document]:
    all_documents = []

    supported_extensions = {
        ".pdf",
        ".txt",
        ".docx",
        ".csv"
    }

    for file_path in folder_path.iterdir():
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            all_documents.extend(load_single_document(file_path))

    return all_documents